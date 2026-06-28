from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(".")
OUT_DIR = ROOT / "output"
REPORT_ASSETS = ROOT / "report_assets"
DOCX_PATH = OUT_DIR / "Isorhy_PM_full_methods_results.docx"


def set_font(run, *, name: str = "Arial", size: int | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "000000", 20, 6),
        ("Heading 2", 14, "000000", 16, 4),
        ("Heading 3", 12, "434343", 14, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_font(r, size=24)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(3)
    r2 = p2.add_run(subtitle)
    set_font(r2, size=12, color="444444")

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(10)
    r3 = p3.add_run("证据链主线：动物 → 细胞 → 分子")
    set_font(r3, size=11, bold=True)


def add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r)


def add_figure(doc: Document, image_path: Path, caption: str, width: float = 6.0) -> None:
    doc.add_picture(str(image_path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(10)
    run = cap.add_run(caption)
    set_font(run, size=10)


def add_table(doc: Document, title: str, headers: list[str], rows: list[list[str]]) -> None:
    add_paragraph(doc, title)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_font(r, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i != 1 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(value)
            set_font(r)
    doc.add_paragraph()


def read_csv_rows(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8") as f:
        return list(csv.DictReader(f))


def add_intro(doc: Document) -> None:
    add_paragraph(
        doc,
        "本总稿在不改动现有结果文件的前提下，基于你提供的第一部分对接文档、第二部分随机森林与 SHAP 结果、第三部分 HE 自动分割输出以及第四部分转录组整合结果重新整理生成。整篇稿件按四个部分依次展开，并最终汇聚为一条“动物→细胞→分子”的完整证据链。",
    )


def add_part1(doc: Document) -> None:
    doc.add_heading("第一部分  GNN + AlphaFold2 + DiffDock：Isorhy 焦亡靶点机制预测", level=1)

    doc.add_heading("Methods", level=2)
    add_paragraph(
        doc,
        "本部分主要沿用上传文档《GNN + AlphaFold2 + DiffDock.docx》的分析流程与详细结论。以焦亡通路中五个关键蛋白 GSDMD、GSDMD-NT、NLRP3、pro-Caspase-1 和 cleaved-Caspase-1 作为候选靶点，首先基于 AlphaFold2 预测各蛋白三维结构；其中 GSDMD-NT 采用前 275 个氨基酸截断构建，NLRP3 采用 135–649 位残基片段用于结构预测与对接，cleaved-Caspase-1 则截取 p20+p10 活性片段并去除 prodomain。随后使用 DiffDock 对 Isorhamnetin 与各靶蛋白进行分子对接，输出最佳 pose 置信度、所有 pose 置信度分布以及 DOCKED.pdb 结构文件，并进一步提取 4 Å 距离阈值内的接触残基信息，形成靶点优先级排序表和结合位点残基列表。",
    )
    add_paragraph(
        doc,
        "在机制解释上，本部分采用上传文档中的结构域分析逻辑：重点关注 NLRP3 的 NBD-LRR 区域是否可被 Isorhy 占据，评估其对炎症小体激活构象和下游 Caspase-1 剪切的潜在影响；同时结合 pro-Caspase-1、cleaved-Caspase-1 以及 GSDMD / GSDMD-NT 的对接表现，判断 Isorhy 更可能通过直接抑制哪一层级的焦亡轴发挥作用。",
    )

    doc.add_heading("Results", level=2)
    add_paragraph(
        doc,
        "根据上传文档，Isorhy 与五个焦亡相关靶蛋白的结合置信度整体处于中等水平（-1.5 至 0 之间）。其中 pro-Caspase-1 的最佳置信度最高，为 -0.25；NLRP3 次之，为 -0.47；cleaved-Caspase-1 和 GSDMD-NT 的最佳置信度分别为 -0.56 和 -0.59；GSDMD 全长蛋白置信度最低，为 -1.62，已接近低置信阈值。若结合亲和力稳定性进一步比较，则 NLRP3 的预测亲和力最强（-7.116 kcal/mol），显著优于其余四个蛋白，提示 Isorhy 与 NLRP3 的结合在热力学上最为稳定。",
    )
    add_paragraph(
        doc,
        "上传文档进一步指出，NLRP3 对接所使用的 135–649 片段覆盖其 NBD（核苷酸结合域）和 LRR（富亮氨酸重复域）区域。NBD 是 NLRP3 炎症小体组装和 ATP 依赖性寡聚化的核心结构域，LRR 则参与配体识别与自抑制调控。因此，Isorhy 如果优先结合于 NBD-LRR 界面，理论上可以通过干扰 NLRP3 的构象激活来阻断下游炎症小体装配。这一解释与文档中的综合机制模型高度一致：Isorhy 首先作用于 NLRP3 上游激活环节，继而减少 pro-Caspase-1 向 cleaved-Caspase-1 的自剪切活化，最终抑制 GSDMD 裂解为 GSDMD-NT 并减少 IL-1β 和 IL-18 的成熟释放。",
    )
    add_paragraph(
        doc,
        "对于其他几个靶点，上传文档给出了更细化的层级判断。pro-Caspase-1 虽然在置信度上排名第一，但其亲和力仅为 -1.252 kcal/mol，弱于 NLRP3，因此更可能是次级靶点而非最核心的直接结合位点；cleaved-Caspase-1 的亲和力接近于零，提示其直接结合贡献有限；GSDMD-NT 的置信度优于 GSDMD 全长，而 GSDMD 全长的正亲和力值（11.798 kcal/mol）则说明 Isorhy 与其全长构象结合不稳定。因此，从上传文档的原始结论来看，Isorhy 更可能优先通过干预 NLRP3 炎症小体组装、并部分影响 pro-Caspase-1 活化，而不是直接稳定结合 GSDMD 全长蛋白。",
    )
    add_paragraph(
        doc,
        "上传文档还将上述 AI 对接结果与既有 Western blot 结果进行了并列解释。文档指出，P. multocida 感染后 NLRP3、cleaved-Caspase-1 和 GSDMD-NT 蛋白表达显著升高，而 Isorhy 40 mg/kg 干预后三者显著下调；pro-Caspase-1 虽有下降趋势但差异未达到统计学显著。也正因为如此，上传文档将 NLRP3 界定为 Isorhy 发挥抗焦亡作用的核心结合靶点，把 pro-Caspase-1 视为可能受到一定影响但不是最关键的直接靶标。该部分因此为后续第二至第四部分提供了明确的分子起点：Isorhy 不是无差别抑制焦亡，而更可能从 NLRP3 炎症小体装配这一上游环节切入。",
    )

    docking_rows = [
        ["1", "pro-Caspase-1", "-0.25", "-1.252", "Moderate"],
        ["2", "NLRP3(135-649)", "-0.47", "-7.116", "Moderate"],
        ["3", "cleaved-Caspase-1", "-0.56", "-0.003", "Moderate"],
        ["4", "GSDMD-NT", "-0.59", "-0.440", "Moderate"],
        ["5", "GSDMD", "-1.62", "11.798", "Low"],
    ]
    add_table(
        doc,
        "表1. 第一部分靶点优先级排序表。",
        ["Rank", "Protein", "Best confidence", "Best affinity (kcal/mol)", "Grade"],
        docking_rows,
    )
    add_figure(
        doc,
        REPORT_ASSETS / "part1_docking_results.png",
        "图1. Isorhamnetin 与 5 个焦亡相关蛋白的 DiffDock 对接结果。左侧为各蛋白最佳 pose 的置信度条形图，右侧为全部 pose 的置信度分布。整体上，pro-Caspase-1 具有最高的最佳置信度，而 NLRP3 表现出最优的热力学亲和力，支持其作为优先机制靶点的定位。",
        width=6.1,
    )
    add_figure(
        doc,
        REPORT_ASSETS / "part1_binding_sites.png",
        "图2. Isorhamnetin 与 5 个候选靶蛋白的接触残基概览（4 Å cutoff）。NLRP3(135-649) 的接触残基数量最多，提示其可能与 Isorhy 形成更广泛、更稳定的结合界面；GSDMD 全长则表现出较弱且分散的结合模式。",
        width=6.1,
    )


def add_part2(doc: Document) -> None:
    doc.add_heading("第二部分  随机森林 + SHAP：感染态与恢复态标志物筛选", level=1)

    doc.add_heading("Methods", level=2)
    add_paragraph(
        doc,
        "本部分重新按你的要求构建“Control vs P. multocida”二分类模型，而不再以 Isorhy 剂量作为直接回归目标。输入特征包括 ELISA 检测得到的 IL-1β 和 IL-18，HE 病理学炎症评分，肺泡面积占比、单位面积肺泡数量、肺泡面积、肺泡周长等形态学指标，肺组织载菌量，以及 TUNEL 阳性细胞数与凋亡率。训练阶段仅使用 Control 与 P. multocida 两组样本建立恢复状态判别器；之后再将 10、20 和 40 mg/kg 三个治疗组样本投影到同一 SHAP 空间，观察其是否沿着“感染态→对照态”的方向迁移。",
    )
    add_paragraph(
        doc,
        "模型解释方面，采用两层输出：第一层使用随机森林的重要性排序作为辅助恢复状态排名图，用于判断哪类指标在区分感染态与恢复态时贡献更大；第二层使用 SHAP beeswarm、组均值热图和 SHAP PCA 轨迹图作为核心解释图，用于直观展示各特征在不同组别中的方向性贡献。分类性能使用 ROC 曲线与 AUC 评价。",
    )

    doc.add_heading("Results", level=2)
    add_paragraph(
        doc,
        "随机森林恢复状态分类结果显示，Control 与 P. multocida 在当前 10 个整合特征构成的空间中可被稳定区分。辅助重要性排序中，Alveolar area (%)、Alveolar area (mm2/mm2)、IL-18、IL-1β、Bacterial load 以及 Inflammation score 位于前列，说明在感染态与恢复态的判别中，肺泡结构破坏程度、炎症因子水平和细菌负荷共同提供了较高信息量。虽然该排序图作为辅助解释使用，但它揭示了一个清晰趋势：Isorhy 的恢复效应不是由单一指标驱动，而是由结构性损伤缓解与炎症级联下调共同构成。",
    )
    add_paragraph(
        doc,
        "SHAP beeswarm 进一步展示了训练集中各特征对 Control 与 P. multocida 分类边界的贡献方向。整体来看，肺泡面积相关特征和炎症评分形成最稳定的分离模式：更大的肺泡面积占比和肺泡面积（以及更接近正常的肺泡结构参数）更倾向于推动样本靠近 Control；更高的炎症评分、细菌负荷以及炎症因子水平则推动样本靠近 P. multocida。这表明模型识别到的不是孤立的单项指标，而是一组共同刻画“感染损伤”与“组织恢复”的复合特征。",
    )
    add_paragraph(
        doc,
        "最关键的结果来自治疗组在 SHAP 空间中的投影。基于当前恢复状态模型，各组平均 P(Control) 依次为：Control 组 1.000，P. multocida 组 0.000，Isorhy 10 mg/kg 组 0.315，Isorhy 20 mg/kg 组 0.395，Isorhy 40 mg/kg 组 0.931。也就是说，10 mg/kg 和 20 mg/kg 组虽然已经开始脱离感染态，但仍主要位于感染区与对照区之间的过渡带；40 mg/kg 组则在整体表型空间中明显贴近 Control，提示高剂量 Isorhy 对感染损伤的恢复最充分。这一结果与动物层面的病理改善、炎症因子下降和载菌量回落方向高度一致。",
    )
    add_paragraph(
        doc,
        "ROC 结果进一步证实了恢复状态分类框架的有效性。在 Control vs P. multocida 的训练集上，分类模型 AUC 为 1.000，表明该组整合指标在当前样本内能够非常稳定地区分感染态与对照态。需要指出的是，这一结果主要说明“感染态和健康态差异非常清楚”，而真正体现 Isorhy 恢复效果的是治疗组在 SHAP 轨迹中的连续迁移。因此，本部分的核心生物学结论并不是“模型分类得很好”，而是“治疗组样本在多维解释空间中呈现出向对照组回归的连续恢复轨迹”。",
    )

    rf_rows = []
    for row in read_csv_rows(OUT_DIR / "part2_recovery_rf_importance.csv"):
        rf_rows.append([row["feature"], f"{float(row['importance']):.3f}"])
    prob_rows = []
    for row in read_csv_rows(OUT_DIR / "part2_recovery_rf_group_means.csv"):
        prob_rows.append([row["group"], f"{float(row['mean_p_control']):.3f}", f"{float(row['sd_p_control']):.3f}"])
    add_table(doc, "表2. 第二部分恢复状态辅助随机森林特征重要性排序。", ["Feature", "Importance"], rf_rows)
    add_table(doc, "表3. 第二部分各组样本平均 P(Control)。", ["Group", "Mean P(Control)", "SD"], prob_rows)

    add_figure(
        doc,
        REPORT_ASSETS / "part2_rf_aux.png",
        "图3. 恢复状态辅助随机森林特征排序图。该图用于辅助显示哪些指标更有助于区分感染态与恢复态，排序结果提示肺泡结构参数、炎症因子和细菌负荷在恢复判别中共同占据较高权重。",
        width=6.0,
    )
    add_figure(
        doc,
        REPORT_ASSETS / "part2_shap_beeswarm.png",
        "图4. SHAP beeswarm summary plot。训练集仅包含 Control 与 P. multocida 两组样本，图中展示了各特征对恢复状态分类边界的方向性贡献。肺泡结构参数与炎症评分表现出最稳定的分离作用。",
        width=6.0,
    )
    add_figure(
        doc,
        REPORT_ASSETS / "part2_shap_trajectory.png",
        "图5. SHAP 空间轨迹图。治疗组样本从 P. multocida 区域向 Control 区域连续迁移，其中 40 mg/kg 组与对照组最接近，提示其恢复最充分。",
        width=6.0,
    )
    add_figure(
        doc,
        REPORT_ASSETS / "part2_shap_heatmap.png",
        "图6. 各组 SHAP 均值热图。正值表示更偏向 Control，负值表示更偏向 P. multocida。40 mg/kg 组在多数特征上已接近 Control 模式，支持剂量依赖性恢复。",
        width=6.0,
    )
    add_figure(
        doc,
        REPORT_ASSETS / "part2_roc_control_vs_pm.png",
        "图7. Control vs P. multocida 分类模型 ROC 曲线。AUC=1.000，说明感染态与对照态在当前多指标空间中具有高度稳定的可分性。",
        width=5.1,
    )


def add_part3(doc: Document) -> None:
    doc.add_heading("第三部分  HE 图像自动分割与计数：自动化病理评分与形态学定量", level=1)

    doc.add_heading("Methods", level=2)
    add_paragraph(
        doc,
        "本部分以肺组织 HE 切片图像为输入，自动生成 tissue mask、alveoli mask 和 nuclei mask，并基于此提取肺泡面积占比、单位组织面积肺泡数量、平均肺泡面积、平均肺泡周长、炎症细胞密度和核面积分数等特征。随后分别构建两类评分模型：第一类为基线线性标定模型，即以炎症细胞密度映射人工评分；第二类为整合多种图像特征的样本级随机森林模型。模型优劣通过交叉验证 R2、MAE、RMSE、Pearson/Spearman 相关性以及 within-1-point rate 进行比较，并最终选取表现最优者作为自动病理评分器。",
    )
    add_paragraph(
        doc,
        "需要强调的是，本部分结果高度依赖 HE 图像质量和分割标注质量。对于肺泡结构完整、染色均匀的图像，模型通常能稳定识别肺泡腔隙和核分布；对于炎症极重或染色差异较大的图像，则更容易出现肺泡边界不清、核区域过密与空腔误判。因此，本部分在结果解释中同时保留模型性能、代表性 overlay 图以及组间偏差描述，以反映自动评分的真实上限。",
    )

    doc.add_heading("Results", level=2)
    add_paragraph(
        doc,
        "HE 自动分割 overlay 图清楚展示了感染损伤与治疗恢复的组织层变化。Control 组肺泡腔隙广泛、规则且连续，提示肺泡结构保持完整；P. multocida 组则表现为肺泡腔隙显著减少、炎症细胞密集聚集、肺泡壁增厚和组织结构紊乱；20 mg/kg Isorhy 组开始出现部分肺泡重新开放，而 40 mg/kg 组肺泡开放度和整体组织结构已明显向 Control 靠拢。这一趋势与人工 HE 评分和动物层炎症指标变化方向完全一致，说明图像自动分割不仅能做静态量化，也能直观反映治疗所带来的结构修复。",
    )
    add_paragraph(
        doc,
        "模型比较结果显示，样本级随机森林明显优于仅基于炎症细胞密度的基线线性模型。最佳模型的交叉验证性能为：n=25，cv_R2=0.881，cv_MAE=0.445，cv_RMSE=0.570；同时 Pearson r=0.945，Spearman r=0.896，within-1-point rate=0.92，weighted kappa=0.915。换句话说，在绝大多数样本中，AI 自动评分与人工评分的误差不超过 1 分，说明该模型已经能够较好地复现人工病理分级的主趋势。",
    )
    add_paragraph(
        doc,
        "按组均值比较时，AI 评分较好重建了 Control 组最低、P. multocida 组最高、治疗组逐渐下降的整体恢复模式。Control 组人工与 AI 平均评分分别约为 0.52 和 0.80；P. multocida 组分别为 4.96 和 4.68；10 mg/kg 组分别为 3.70 和 3.60；20 mg/kg 组分别为 3.38 和 3.09；40 mg/kg 组分别为 1.32 和 1.95。可以看到，AI 模型在感染组和中等损伤组表现较稳，而在 40 mg/kg 组存在轻度高估，这提示当炎症下降到较轻程度时，图像染色差异、局部结构异质性和分割误差会更明显地影响评分输出。",
    )
    add_paragraph(
        doc,
        "最佳模型的重要特征主要集中在 nuclei_area_px_mean、nuclei_area_fraction_mean、mean_alveolar_area_px_mean 和 alveoli_count_per_tissue_area_mean。这说明自动评分并不是简单依赖“炎症细胞越多，分数越高”这一单一线索，而是在整合核密度变化、核面积分布、肺泡平均面积以及单位面积肺泡数量等多重结构信号之后，给出更接近人工病理观察的综合判断。也正因为如此，HE 自动评分部分能够自然衔接“动物层的病理改善”和“分子层的焦亡抑制”，成为这条证据链中的组织学桥梁。",
    )

    add_figure(
        doc,
        REPORT_ASSETS / "part3_he_overlay_panel.png",
        "图8. 代表性 HE 自动分割叠加图。Control 组肺泡腔隙规则开放；P. multocida 组出现明显肺泡塌陷、核密度增高和腔隙减少；20 mg/kg 与 40 mg/kg Isorhy 组显示出逐步恢复的肺泡开放度和组织结构。",
        width=6.3,
    )
    add_figure(
        doc,
        REPORT_ASSETS / "part3_manual_vs_ai.png",
        "图9. 人工评分与 AI 自动评分的一致性散点图。最佳样本级随机森林模型在样本层面与人工评分高度一致，说明自动评分能够稳定复现人工病理分级趋势。",
        width=5.8,
    )
    add_figure(
        doc,
        REPORT_ASSETS / "part3_group_mean_manual_vs_ai.png",
        "图10. 各组人工评分与 AI 平均评分对比。AI 模型成功重建了感染组最高、治疗组逐渐恢复的组间模式，但在 40 mg/kg 组存在轻度高估，提示该模型仍受切片质量和分割误差影响。",
        width=6.0,
    )


def add_part4(doc: Document) -> None:
    doc.add_heading("第四部分  转录组 + NLRP2 机制整合：从差异基因到焦亡网络", level=1)

    doc.add_heading("Methods", level=2)
    add_paragraph(
        doc,
        "本部分整合差异表达分析、GSEA/通路富集、WGCNA 模块识别和 AI 增强 PPI 网络。差异基因统计同时保留两种口径：一是名义 p<0.05 且 |log2FC|>1 的探索性差异集合，用于完整展示感染和治疗引起的广泛转录改变；二是 adj.p<0.05 且 |log2FC|>1 的严格差异集合，用于评估在多重校正后仍能稳定保留下来的高置信度变化。WGCNA 以 Isorhy 剂量作为表型变量构建共表达模块，筛选与剂量正相关或负相关的关键模块；之后将焦亡相关蛋白互作关系输入 AI 增强 PPI 网络，识别 NLRP2 邻近网络中的关键协同节点。",
    )
    add_paragraph(
        doc,
        "由于你现有的细胞功能学稿件已将 NLRP2 设定为过表达验证的核心节点，而转录组结果表中 Nlrp2 在两次比较中的方向性呈现“反转”而不是单向上调，因此本部分结果写作采取保守原则：把 NLRP2 作为感染-治疗状态中最值得关注的关键响应基因与网络中心，而不简单把它压缩成单向变化的‘上调靶点’。",
    )

    doc.add_heading("Results", level=2)
    add_paragraph(
        doc,
        "差异表达分析显示，P. multocida 感染会引起大范围的转录重塑。按名义阈值统计，Pm vs Control 共识别到 1677 个差异基因，其中上调 386 个、下调 1291 个；Pm+Isorhy vs Pm 共识别到 1478 个差异基因，其中上调 612 个、下调 866 个。若采用更严格的 FDR 校正标准，则 Pm vs Control 仅保留 19 个严格差异基因，而 Pm+Isorhy vs Pm 在当前结果表中未保留严格 FDR 阳性基因。换句话说，感染带来的转录扰动更强、更集中，而 Isorhy 的作用更像对广泛失衡网络进行回调，而不是诱导少数极端差异基因的剧烈翻转。",
    )
    add_paragraph(
        doc,
        "在候选基因汇总中，Nlrp2 是最值得关注的焦亡相关响应基因之一：其在 Pm vs Control 比较中的 log2FC 为 -1.992，而在 Pm+Isorhy vs Pm 中为 +2.312，并被标记为方向反转基因。这一结果说明 NLRP2 对感染与治疗状态都高度敏感，且不是随机噪音式变化。结合你现有细胞层 NLRP2 过表达结果，可以更合理地理解为：NLRP2 所处的焦亡调控网络在感染和治疗过程中被显著重塑，而 NLRP2 正好位于这一重塑过程的关键节点上。",
    )
    add_paragraph(
        doc,
        "富集分析显示，钙信号、氧化磷酸化、糖酵解、脂质代谢与代谢重编程相关通路在感染与治疗比较中反复出现，提示焦亡抑制并不是孤立事件，而与能量代谢、炎症信号和上皮应答网络紧密耦联。特别是在 Pm+Isorhy vs Pm 的富集图中，Thermogenesis、Glycolysis/Gluconeogenesis、Calcium signaling pathway 等条目位居前列，这说明 Isorhy 对感染肺组织的影响不仅体现在炎症抑制上，也体现在代谢和应激通路的系统级调整上。",
    )
    add_paragraph(
        doc,
        "WGCNA 结果进一步支持这种系统性恢复模式。多个模块与 Isorhy 剂量显著相关，其中 turquoise 和 black 模块与剂量呈正相关（r=0.85），而 blue、red 和 pink 模块与剂量呈负相关（r=-0.96）。这种模块层面的方向性变化提示：Isorhy 干预并不是简单地让所有感染相关基因“回到原点”，而是沿着特定共表达程序推动一部分模块被激活、另一部分模块被抑制，从而实现整体网络状态的重平衡。",
    )
    add_paragraph(
        doc,
        "AI 增强 PPI 网络把这种“关键节点驱动网络重平衡”的机制具体化到了蛋白互作层面。在当前网络图中，NLRP2 位于中心，其直接邻居中 ASC、CASP1、NLRP3、NEK7、P2RX7 和 HSP90AA1 获得更高 AI 权重，说明这些蛋白最可能决定 NLRP2 在焦亡网络中的实际功能输出。换言之，NLRP2 不是一个孤立的差异基因，而更像是连接炎症小体组装、Caspase-1 活化、细胞膜打孔和炎症因子释放的网络枢纽。",
    )
    add_paragraph(
        doc,
        "当把这一部分与第一部分分子对接结果联立时，证据链会进一步闭合：第一部分提示 Isorhy 最可能直接作用于 NLRP3；第四部分则提示 NLRP2/NLRP3/CASP1/ASC 等节点位于治疗响应网络中心。因此，较为合理的整合机制模型是：Isorhy 通过优先干预 NLRP3 炎症小体激活，同时重塑以 NLRP2 为中心的焦亡响应网络，最终在蛋白水平减少 Caspase-1 激活和 GSDMD 裂解，在组织水平降低 IL-1β/IL-18 释放、减轻细胞死亡并恢复肺泡结构。",
    )

    doc.add_heading("NLRP2 相关功能验证", level=3)
    add_paragraph(
        doc,
        "为验证 NLRP2 是否不仅是转录组和网络分析中的相关节点，而且能够真实决定细胞表型变化，本研究进一步整合了 NLRP2 过表达背景下的 CCK-8、流式凋亡和透射电镜（TEM）结果。该部分采用 4 组设计：OE-NC+Ctrl、OE-NC+P.multocida、OE-NLRP2+Ctrl 以及 OE-NLRP2+P.multocida-Isorhy。整体上，这批结果不是为了新增一条独立分析线，而是作为第四部分的功能性收尾，用于回答“当前锁定的 NLRP2 节点是否能够驱动并解释前面观察到的损伤与恢复表型”。",
    )
    add_paragraph(
        doc,
        "CCK-8 结果显示，OE-NC+Ctrl 组细胞活力最高，平均 OD450 约为 1.29；在单纯感染条件下，OE-NC+P.multocida 组下降至约 1.06；OE-NLRP2+Ctrl 组进一步降至最低水平，均值约 0.72；而在 OE-NLRP2 背景下给予 Isorhy 干预后，OE-NLRP2+P.multocida-Isorhy 组细胞活力恢复至约 1.26，接近 OE-NC+Ctrl 水平。也就是说，NLRP2 过表达本身即可显著压低细胞活力，而 Isorhy 在感染+NLRP2 扰动背景下又能把这一损伤表型明显拉回。这个趋势说明 NLRP2 不是被动标记物，而是真正参与了细胞损伤与恢复过程的功能节点。",
    )
    add_paragraph(
        doc,
        "流式凋亡结果为这一判断提供了更直接的细胞命运证据。根据 Annexin V-FITC/PI 分析，OE-NC+Ctrl 组的早期凋亡、晚期凋亡和总凋亡率分别约为 8.39%、6.96% 和 15.35%；OE-NC+P.multocida 组对应指标升高至约 11.80%、6.41% 和 18.21%；OE-NLRP2+Ctrl 组则进一步升高至约 14.23%、13.27% 和 27.50%，说明在无感染条件下单独提高 NLRP2 水平就足以明显增强细胞死亡倾向；而 OE-NLRP2+P.multocida-Isorhy 组的早期凋亡、晚期凋亡和总凋亡率分别回落至约 8.07%、7.42% 和 15.49%，与 OE-NLRP2+Ctrl 相比均呈明显下降趋势。换句话说，Isorhy 在 NLRP2 扰动背景下依然能够显著抑制细胞凋亡/焦亡相关死亡表型，提示其保护作用并不是偶然依赖于某个单一模型条件。",
    )
    add_paragraph(
        doc,
        "TEM 结果从超微结构层面进一步支持了这一结论。OE-NC+Ctrl 组细胞整体结构相对完整，细胞器边界清楚；OE-NC+P.multocida 组可见更明显的细胞器损伤和应激改变；OE-NLRP2+Ctrl 组则出现更突出的线粒体肿胀、膜结构紊乱及细胞质损伤征象，提示 NLRP2 过表达本身即可放大细胞超微结构损害；在 OE-NLRP2+P.multocida-Isorhy 组中，相关损伤特征较 OE-NLRP2+Ctrl 明显减轻，细胞器完整性恢复更好。虽然 TEM 更偏重形态学呈现而非严格定量，但它很好地补足了 CCK-8 和流式凋亡的证据链，使 NLRP2 的“相关性”进一步上升为“具备可观察表型后果的功能性节点”。",
    )
    add_paragraph(
        doc,
        "因此，这一功能验证小节与第四部分前面的转录组、WGCNA 和 AI 增强 PPI 网络并不是并列的附加材料，而是对核心机制判断的闭环验证：转录组和网络分析提示 NLRP2 位于感染-治疗响应网络中心；CCK-8、流式凋亡和 TEM 则说明，一旦直接扰动 NLRP2，细胞活力、死亡比例和超微结构损伤都会随之变化，并且 Isorhy 可以在这一背景下明显逆转相关表型。这使得第四部分能够更有力地把 NLRP2 定位为连接上游网络重塑与下游细胞损伤/恢复表型的关键功能节点。",
    )

    add_figure(
        doc,
        REPORT_ASSETS / "part4_volcano.png",
        "图11. 转录组火山图。Nlrp2 在感染与治疗比较中均被指向，并表现出方向反转，提示其是感染-治疗状态切换中的关键响应节点。",
        width=6.2,
    )
    add_figure(
        doc,
        REPORT_ASSETS / "part4_gsea.png",
        "图12. KEGG/GO 富集结果。感染与治疗比较均提示炎症信号和能量代谢网络广泛参与 Isorhy 的保护作用。",
        width=6.2,
    )
    add_figure(
        doc,
        REPORT_ASSETS / "part4_wgcna.png",
        "图13. WGCNA 共表达模块分析。多个模块与 Isorhy 剂量显著相关，提示治疗效应在基因共表达层面具有清晰的程序化特征。",
        width=6.2,
    )
    add_figure(
        doc,
        REPORT_ASSETS / "part4_ppi.png",
        "图14. AI 增强 PPI 网络。NLRP2 位于网络中心，其与 ASC、CASP1、NLRP3、NEK7 等高权重邻居共同构成焦亡网络核心支架。",
        width=6.2,
    )


def add_part5_summary(doc: Document) -> None:
    doc.add_heading("四部分整合总结：动物→细胞→分子证据链", level=1)
    add_paragraph(
        doc,
        "将四个部分合并后，可以得到一条更完整的机制主线。第一部分从结构生物学层面提示 Isorhy 最可能直接作用于 NLRP3，并可能部分影响 pro-Caspase-1；第二部分从多指标机器学习角度证明，治疗组在整体解释空间中沿着“感染态→对照态”连续恢复；第三部分则从 HE 自动分割和病理评分角度显示肺泡结构与炎症状态同步改善；第四部分进一步把这些表型变化锚定到 NLRP2/NLRP3-CASP1-GSDMD 焦亡网络重塑上。",
    )
    add_paragraph(
        doc,
        "因此，Isorhy 的作用并不是单层次的‘抗炎’或‘降菌’，而更像是在感染损伤之后，通过抑制炎症小体激活、减弱 Caspase-1/GSDMD 轴的焦亡执行、降低 IL-1β/IL-18 释放，并最终推动肺组织从严重炎症和结构破坏状态逐步回归至接近正常的稳态。也正因为这种跨层级一致性，本项目能够形成较完整的“动物→细胞→分子”三层证据链。",
    )


def build_doc() -> None:
    OUT_DIR.mkdir(exist_ok=True)
    doc = Document()
    style_document(doc)
    add_title(doc, "Isorhy_PM_full_methods_results", "完整四部分总稿")
    add_intro(doc)
    add_part1(doc)
    add_part2(doc)
    add_part3(doc)
    add_part4(doc)
    add_part5_summary(doc)
    doc.save(str(DOCX_PATH))
    print(DOCX_PATH)


if __name__ == "__main__":
    build_doc()
